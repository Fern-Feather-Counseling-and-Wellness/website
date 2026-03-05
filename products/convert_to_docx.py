#!/usr/bin/env python3
"""Convert attachment workbook markdown to Word document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_heading(doc, text, level=1):
    """Add a heading with consistent styling."""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(34, 74, 54)  # Dark green
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(34, 74, 54)
    else:
        run.font.size = Pt(12)
        run.font.bold = True
    return heading

def add_paragraph(doc, text, style=None):
    """Add a paragraph with optional style."""
    p = doc.add_paragraph(style=style)
    
    # Handle inline bold (text surrounded by **)
    parts = text.split('**')
    for i, part in enumerate(parts):
        run = p.add_run(part)
        if i % 2 == 1:  # Odd indices are bold
            run.bold = True
    
    return p

def add_note_box(doc, text):
    """Add a note/callout box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(f"🧠 {text}")
    run.italic = True
    run.font.color.rgb = RGBColor(60, 60, 60)

def add_quote(doc, text):
    """Add a quote/affirmation."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    run = p.add_run(f"\"{text}\"")
    run.italic = True

def main():
    doc = Document()
    
    # Set document title
    title = doc.add_heading(level=0)
    title_run = title.add_run("ATTACHMENT BIOLOGY WORKBOOK")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(34, 74, 54)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph("A Somatic Toolkit for Understanding (& Rewiring) Your Relational Wiring")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(80, 80, 80)
    
    doc.add_paragraph()
    
    # Welcome section
    add_heading(doc, "Welcome", level=1)
    add_paragraph(doc, "Traditional attachment theory tells us we have \"styles.\" But you are not a label — you are a living, adaptable system.")
    add_paragraph(doc, "This workbook is built on the simple, powerful idea: **Attachment isn't personality or pathology — it's biology. What was wired can be rewired.**")
    add_paragraph(doc, "Inside, you'll learn to recognize your patterns, trace their wiring, and shift them in real time using somatic practices, reflective exercises, and neuroscience-backed strategies.")
    add_paragraph(doc, "This is science meeting self-compassion. Let's begin rewiring your relational biology — together.")
    
    # Disclaimer
    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer_run = disclaimer.add_run("⚠️ Disclaimer: This workbook is for educational and self-reflection purposes only. It is not a substitute for professional mental health treatment. If you are experiencing significant distress, please seek support from a licensed therapist.")
    disclaimer_run.font.size = Pt(9)
    disclaimer_run.font.color.rgb = RGBColor(120, 120, 120)
    
    # PART I
    doc.add_page_break()
    part1_title = doc.add_heading(level=1)
    part1_run = part1_title.add_run("PART I: FOUNDATIONS")
    part1_run.font.size = Pt(20)
    part1_run.font.color.rgb = RGBColor(34, 74, 54)
    
    add_paragraph(doc, "What Is Attachment Biology?", style='Heading 2')
    add_heading(doc, "Beyond Styles to Patterns", level=2)
    add_paragraph(doc, "Attachment theory describes how our earliest relationships shape our expectations about connection, safety, and emotional needs.")
    add_paragraph(doc, "But attachment is not fixed personality. It's **predictive patterning** — your nervous system learned patterns to help you survive and belong. These patterns are adaptive, not defective. They're biology responding to environment.")
    
    add_note_box(doc, "Your attachment pattern is a survival strategy your nervous system developed — not a flaw, but protective wiring.")
    
    add_heading(doc, "The Neuroscience in Plain Language", level=3)
    add_paragraph(doc, "• Our brains are prediction machines")
    add_paragraph(doc, "• Early experiences shape our \"internal working model\" of relationships")
    add_paragraph(doc, "• This model lives in implicit memory — body-based, below conscious awareness")
    add_paragraph(doc, "• We don't just think our attachment patterns — we feel, sense, and react them")
    
    add_heading(doc, "The Three Biological Grounds of Attachment", level=3)
    add_paragraph(doc, "Protest — The drive to seek connection, often anxiously")
    add_paragraph(doc, "Despair — The shut-down when connection feels unavailable")
    add_paragraph(doc, "Detachment — The \" functional freeze\" protecting us from further hurt")
    
    add_quote(doc, "I am learning to trust the wisdom of my body, even as I gently challenge its old predictions.")
    
    # PART II
    doc.add_page_break()
    part2_title = doc.add_heading(level=1)
    part2_run = part2_title.add_run("PART II: MAPPING YOUR BIOLOGY")
    part2_run.font.size = Pt(20)
    part2_run.font.color.rgb = RGBColor(34, 74, 54)
    
    add_heading(doc, "The Attachment Biology Framework", level=2)
    add_paragraph(doc, "We are mapping the three processes that create attachment patterns:")
    add_paragraph(doc, "• Activation — How your system revs up, escalates, or clings")
    add_paragraph(doc, "• Regulation — How you return to baseline, settle, restore")
    add_paragraph(doc, "• Orientation — Where your attention goes — toward or away from connection")
    
    add_heading(doc, "Module 1: Your Activation Signature", level=2)
    add_heading(doc, "Understanding Your Biological Response", level=3)
    add_paragraph(doc, "Describe what happens in your body when connection feels threatened or uncertain:")
    
    # Journal prompts
    prompts = [
        ("Where do you feel it?", "(chest tightness, stomach churning, throat constriction, etc.)"),
        ("What sensations arise?", "(heat, cold, tingling, numbness, racing heart)"),
        ("What impulses emerge?", "(reach out, withdraw, freeze, get busy, etc.)"),
        ("What thoughts typically accompany this activation?", "")
    ]
    
    for prompt, hint in prompts:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run(f"• {prompt}")
        run.bold = True
        if hint:
            p.add_run(f" {hint}")
        # Add space for writing
        space = doc.add_paragraph()
        for _ in range(4):
            space.add_run("_" * 80 + "\n")
        space.paragraph_format.space_after = Pt(12)
    
    add_heading(doc, "Module 2: Your Regulation Patterns", level=2)
    add_paragraph(doc, "Regulation is your capacity to settle, soothe, and return to baseline. It's not suppression — it's completion.")
    
    reg_prompts = [
        "What happens to your body when distress finally passes?",
        "Who or what has helped you feel settled in the past?",
        "What self-soothing strategies do you naturally use?",
        "What makes regulation harder for you?"
    ]
    
    for prompt in reg_prompts:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run(f"• {prompt}")
        run.bold = True
        space = doc.add_paragraph()
        for _ in range(3):
            space.add_run("_" * 80 + "\n")
    
    add_note_box(doc, "Regulation resources are anything that helps you feel safe, settled, and connected — internal and external.")
    
    add_heading(doc, "Module 3: Your Orientation Patterns", level=2)
    add_paragraph(doc, "Orientation is where your attention automatically goes when relationships feel stressful.")
    
    orient_prompts = [
        "When connection feels uncertain, I tend to focus on...",
        "When I'm upset with someone, I typically...",
        "My attention naturally gravitates toward...",
        "I avoid paying attention to..."
    ]
    
    for prompt in orient_prompts:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run(f"• {prompt}")
        run.bold = True
        space = doc.add_paragraph()
        for _ in range(3):
            space.add_run("_" * 80 + "\n")
    
    # PART III
    doc.add_page_break()
    part3_title = doc.add_heading(level=1)
    part3_run = part3_title.add_run("PART III: REWIRING PRACTICES")
    part3_run.font.size = Pt(20)
    part3_run.font.color.rgb = RGBColor(34, 74, 54)
    
    add_paragraph(doc, "Each practice targets a specific biological process: activation, regulation, or orientation. They are cumulative — the more you practice, the more your wiring shifts.")
    
    # Practice 1
    add_heading(doc, "Practice 1: Somatic Grounding for Protest States", level=2)
    add_paragraph(doc, "Best for: Anxious activation, racing thoughts, reaching/seeking urges")
    add_heading(doc, "The Practice", level=3)
    steps = [
        "Place both feet flat on the floor. Feel the weight.",
        "Place one hand on your chest, one on your belly.",
        "Breathe slowly — in through nose for 4 counts, out for 6.",
        "As you exhale, say to yourself: \"I am here. I am safe enough.\"",
        "Notice: What changes in your body? (Don't force — just notice)"
    ]
    for i, step in enumerate(steps, 1):
        add_paragraph(doc, f"{i}. {step}")
    
    add_heading(doc, "The Science", level=3)
    add_paragraph(doc, "This activates your ventral vagal system — the \"safe and social\" branch of your nervous system. Grounding signals to your biology that you are not in immediate danger.")
    
    add_heading(doc, "Reflection", level=3)
    refl1 = [
        "What did I notice in my body during this practice?",
        "How did my activation level change (0-10)?",
        "What was hardest or easiest about staying present?"
    ]
    for prompt in refl1:
        p = doc.add_paragraph()
        run = p.add_run(f"• {prompt}")
        run.bold = True
        space = doc.add_paragraph()
        space.add_run("_" * 80 + "\n" * 3)
    
    # Practice 2
    doc.add_page_break()
    add_heading(doc, "Practice 2: Containment for Despair States", level=2)
    add_paragraph(doc, "Best for: Shut-down, numbness, withdrawal, collapse")
    add_heading(doc, "The Practice", level=3)
    steps2 = [
        "Cross your arms and give yourself a gentle hug (or wrap arms around yourself if preferred).",
        "Rock slightly — back and forth, or side to side.",
        "Make a low sound: \"Mmmm\" or \"Hummmm\" — vibration helps.",
        "Visualize: A warm container around you. Light and safe.",
        "Stay here for 2-3 minutes, noticing without forcing."
    ]
    for i, step in enumerate(steps2, 1):
        add_paragraph(doc, f"{i}. {step}")
    
    add_heading(doc, "The Science", level=3)
    add_paragraph(doc, "Self-soothing touch and rocking activate C-tactile fibers and stimulate vagal tone — literally signaling safety to your nervous system.")
    
    add_heading(doc, "Reflection", level=3)
    for prompt in ["What did I notice about my internal state?", "What sensations emerged, even subtly?", "How does this practice change my sense of safety?"]:
        p = doc.add_paragraph()
        run = p.add_run(f"• {prompt}")
        run.bold = True
        space = doc.add_paragraph()
        space.add_run("_" * 80 + "\n" * 3)
    
    # Practice 3
    doc.add_page_break()
    add_heading(doc, "Practice 3: Orienting for Disorganized/Detached States", level=2)
    add_paragraph(doc, "Best for: Freeze, fog, dissociation, feeling disconnected from self or others")
    add_heading(doc, "The Practice", level=3)
    steps3 = [
        "Sit comfortably. Let your eyes wander the room.",
        "Name 5 things you can see (colors, shapes, textures).",
        "Name 3 things you can hear.",
        "Place one hand on a nearby surface — feel texture, temperature.",
        "Return to your body: What's one sensation you can name right now?"
    ]
    for i, step in enumerate(steps3, 1):
        add_paragraph(doc, f"{i}. {step}")
    
    add_heading(doc, "The Science", level=3)
    add_paragraph(doc, "Orienting engages your nervous system's ability to scan for safety. It brings attention outward then slowly back inward — rebuilding the sense of \"here-ness.\"")
    
    add_heading(doc, "Reflection", level=3)
    for prompt in ["How present do I feel right now (0-10)?", "What brought me back most effectively?", "What felt challenging about staying with my senses?"]:
        p = doc.add_paragraph()
        run = p.add_run(f"• {prompt}")
        run.bold = True
        space = doc.add_paragraph()
        space.add_run("_" * 80 + "\n" * 3)
    
    # Practice 4
    doc.add_page_break()
    add_heading(doc, "Practice 4: Boundary Practice for Orientation Shifts", level=2)
    add_paragraph(doc, "Best for: Over-focusing on others, losing yourself in relationships, anxiety about others' states")
    add_heading(doc, "The Practice", level=3)
    steps4 = [
        "Stand or sit. Feel your edges — where your body ends.",
        "Set a gentle intention: \"I can care about you AND stay with me.\"",
        "Breathe in: \"I am here.\" Breathe out: \"You are there.\"",
        "Notice: What sensations arise when you hold this distinction?",
        "Stay with yourself for 2 minutes — not abandoning self for other."
    ]
    for i, step in enumerate(steps4, 1):
        add_paragraph(doc, f"{i}. {step}")
    
    add_heading(doc, "The Science", level=3)
    add_paragraph(doc, "Healthy boundaries are somatic, not cognitive. This practice trains your nervous system to hold connection AND separateness — the hallmark of secure attachment.")
    
    add_heading(doc, "Reflection", level=3)
    for prompt in ["What did I notice about my capacity to stay with myself?", "What sensations accompanied this boundary?", "What belief about boundaries showed up?"]:
        p = doc.add_paragraph()
        run = p.add_run(f"• {prompt}")
        run.bold = True
        space = doc.add_paragraph()
        space.add_run("_" * 80 + "\n" * 3)
    
    add_note_box(doc, "Reference: Dr. Pat Ogden's Sensorimotor Psychotherapy — boundaries as somatic experience.")
    
    # PART IV
    doc.add_page_break()
    part4_title = doc.add_heading(level=1)
    part4_run = part4_title.add_run("PART IV: RELATIONAL REWIRING")
    part4_run.font.size = Pt(20)
    part4_run.font.color.rgb = RGBColor(34, 74, 54)
    
    add_heading(doc, "Practicing New Patterns with Others", level=2)
    add_paragraph(doc, "Individual practices shift your baseline. But attachment rewires in relationship — with safe others, with yourself-as-other.")
    add_paragraph(doc, "This section focuses on translating your somatic work into real-time relational shifts.")
    
    add_heading(doc, "Exercise: The Pause Practice", level=3)
    add_paragraph(doc, "When you feel triggered in relationship, try this micro-intervention:")
    pause_steps = [
        "PAUSE — Notice the urge to react (reach, withdraw, fix, etc.)",
        "GROUND — One breath. Feet on floor. Hand on heart or belly.",
        "CHOOSE — Ask: What would my wise adult self do here?",
        "ACT — Respond (not react) from this slightly shifted state."
    ]
    for i, step in enumerate(pause_steps, 1):
        add_paragraph(doc, f"{i}. {step}")
    
    add_note_box(doc, "The pause is your superpower. Even one second of awareness interrupts pattern repetition.")
    
    add_heading(doc, "Journal Prompt: After a Triggered Moment", level=3)
    for prompt in [
        "What activated me? (Situation, words, sensation)",
        "What did my old pattern want me to do?",
        "What did I actually do?",
        "What one thing helped me stay present or return to center?",
        "What would I want to try differently next time?"
    ]:
        p = doc.add_paragraph()
        run = p.add_run(f"• {prompt}")
        run.bold = True
        space = doc.add_paragraph()
        space.add_run("_" * 80 + "\n" * 3)
    
    add_heading(doc, "Understanding Your Triggers", level=3)
    add_paragraph(doc, "Complete this map for your most common triggers:")
    
    p = doc.add_paragraph()
    run = p.add_run("Trigger:")
    run.bold = True
    doc.add_paragraph("______________________________________________")
    
    p = doc.add_paragraph()
    run = p.add_run("Body sensation:")
    run.bold = True
    doc.add_paragraph("______________________________________________")
    
    p = doc.add_paragraph()
    run = p.add_run("Old pattern response:")
    run.bold = True
    doc.add_paragraph("______________________________________________")
    
    p = doc.add_paragraph()
    run = p.add_run("New response I want to practice:")
    run.bold = True
    doc.add_paragraph("______________________________________________")
    
    p = doc.add_paragraph()
    run = p.add_run("One micro-skill to remember:")
    run.bold = True
    doc.add_paragraph("______________________________________________")
    
    # PART V
    doc.add_page_break()
    part5_title = doc.add_heading(level=1)
    part5_run = part5_title.add_run("PART V: INTEGRATION")
    part5_run.font.size = Pt(20)
    part5_run.font.color.rgb = RGBColor(34, 74, 54)
    
    add_heading(doc, "Your Personalized Rewiring Plan", level=2)
    add_paragraph(doc, "Use this section to design your ongoing practice. Rewiring happens through repetition — small practices, consistently.")
    
    add_heading(doc, "Daily Grounding Practice", level=3)
    add_paragraph(doc, "Choose one practice from Part III to do daily. Which one resonates most?")
    p = doc.add_paragraph()
    for _ in range(3):
        p.add_run("_" * 80 + "\n")
    
    add_paragraph(doc, "When will you practice? (Morning, before bed, during transitions?)")
    p = doc.add_paragraph()
    for _ in range(2):
        p.add_run("_" * 80 + "\n")
    
    add_heading(doc, "In-the-Moment Resources", level=3)
    add_paragraph(doc, "What are 3 things you can do when activated in real-time?")
    for i in range(1, 4):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}.")
        run.bold = True
        doc.add_paragraph("______________________________________________")
    
    add_heading(doc, "Support Plan", level=3)
    add_paragraph(doc, "Who supports your growth? List supportive people or resources:")
    p = doc.add_paragraph()
    for _ in range(4):
        p.add_run("_" * 80 + "\n")
    add_paragraph(doc, "When would you consider reaching out for support?")
    p = doc.add_paragraph()
    for _ in range(3):
        p.add_run("_" * 80 + "\n")
    
    add_heading(doc, "Progress Tracking", level=3)
    add_paragraph(doc, "Check in weekly. Rate each area 1-10:")
    add_paragraph(doc, "Week of ___________")
    for area in ["Overall sense of safety in body", "Ability to self-regulate", "Capacity to stay present in conflict", "Ease returning to baseline after upset", "Connection to self", "Satisfaction in relationships"]:
        doc.add_paragraph(f"{area}: ____/10")
    
    add_heading(doc, "Celebration & Compassion", level=2)
    add_paragraph(doc, "Growth is not linear. You will have setbacks. This is normal.")
    add_paragraph(doc, "What matters is your commitment to the process — showing up, again and again, with curiosity and kindness toward yourself.")
    add_quote(doc, "I am not fixing myself. I am learning to support the system that has been supporting me all along.")
    
    # APPENDIX
    doc.add_page_break()
    app_title = doc.add_heading(level=1)
    app_run = app_title.add_run("APPENDIX")
    app_run.font.size = Pt(20)
    app_run.font.color.rgb = RGBColor(34, 74, 54)
    
    add_heading(doc, "Quick Reference: The Four Practices", level=2)
    
    ref_data = [
        ("Somatic Grounding", "Anxious activation, racing thoughts", "Feet flat, hand on heart/belly, slow breathing"),
        ("Containment", "Shut-down, numbness, despair", "Self-hug, rocking, humming, visualization"),
        ("Orienting", "Freeze, dissociation, fog", "5-4-3-2-1 senses technique"),
        ("Boundary Practice", "Losing self in others, over-focus on others", "Feel your edges, stay with self while caring")
    ]
    
    # Create table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Practice'
    hdr_cells[1].text = 'Best For'
    hdr_cells[2].text = 'Core Technique'
    
    for name, best_for, technique in ref_data:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = best_for
        row_cells[2].text = technique
    
    add_heading(doc, "Recommended Reading", level=2)
    resources = [
        "Bessel van der Kolk — The Body Keeps the Score",
        "Stephen Porges — The Polyvagal Theory (theory)",
        "Deb Dana — Anchored, Polyvagal Exercises for Safety and Connection",
        "Pat Ogden — Sensorimotor Psychotherapy",
        "Diane Poole Heller — The Power of Attachment",
        "Stan Tatkin — Wired for Love"
    ]
    for resource in resources:
        doc.add_paragraph(f"• {resource}")
    
    add_heading(doc, "When to Seek Professional Support", level=2)
    add_paragraph(doc, "This workbook is a complement to, not replacement for, therapeutic work. Consider professional support if:")
    indicators = [
        "You experience trauma flashbacks or intrusive memories",
        "Your distress significantly interferes with daily functioning",
        "You're experiencing thoughts of self-harm",
        "You're in an unsafe relationship or environment",
        "You've tried these practices and feel stuck or more distressed"
    ]
    for ind in indicators:
        doc.add_paragraph(f"• {ind}")
    
    add_paragraph(doc, "Finding a trauma-informed therapist: Psychology Today therapist directory, search for \"somatic,\" \"attachment-based,\" \"trauma-informed.\"")
    
    # Final page
    doc.add_page_break()
    doc.add_paragraph()
    doc.add_paragraph()
    closing = doc.add_paragraph("Your attachment patterns were adaptive. They kept you safe. They helped you belong.")
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing.runs[0].font.size = Pt(14)
    closing.runs[0].italic = True
    
    doc.add_paragraph()
    closing2 = doc.add_paragraph("Now, with awareness and practice, you can update them — not because you were broken, but because you're growing.")
    closing2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing2.runs[0].font.size = Pt(14)
    closing2.runs[0].italic = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    closing3 = doc.add_paragraph("Science. Soma. Self-Compassion.")
    closing3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing3.runs[0].font.size = Pt(12)
    closing3.runs[0].bold = True
    
    doc.add_paragraph()
    closing4 = doc.add_paragraph("This is Attachment Biology.")
    closing4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing4.runs[0].font.size = Pt(16)
    closing4.runs[0].bold = True
    closing4.runs[0].font.color.rgb = RGBColor(34, 74, 54)
    
    # Save
    output_path = "/home/agent/.openclaw/workspace/products/Attachment-Biology-Workbook.docx"
    doc.save(output_path)
    print(f"✓ Created: {output_path}")

if __name__ == "__main__":
    main()
